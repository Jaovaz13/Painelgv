"""
Geração de relatório em formato Microsoft Word (.docx)
Substitui o formato PDF anterior, mantendo as análises de tendência e dados gerais.
"""
import logging
from pathlib import Path
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

from analytics.tendencias import analisar_tendencia
from analytics.estimativa_pib import get_estimativa_stored
from config import MUNICIPIO, UF, DATA_DIR
from database import get_timeseries, list_indicators
from reports.abnt import apply_abnt_styles, add_figure_caption
from reports.charts import ChartGenerator
from config.indicators import CATALOGO_INDICADORES

logger = logging.getLogger(__name__)

TITULO_SECRETARIA = "Secretaria Municipal de Desenvolvimento, Ciência, Tecnologia e Inovação"

def _fmt_num(val: float) -> str:
    """Formata número para exibição em português (vírgula decimal)."""
    try:
        if val >= 1_000_000:
            return f"{val/1_000_000:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") + " mi"
        return f"{val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return str(val)

def gerar_relatorio_docx(
    ano_inicio: int,
    ano_fim: int,
    output_path: str | Path | None = None,
) -> Path:
    """
    Gera relatório Word com histórico, tendência e dados tabulares.
    """
    output_path = Path(output_path) if output_path else DATA_DIR / f"relatorio_{ano_inicio}_{ano_fim}.docx"
    output_path = output_path.resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    logger.info("Gerando relatório DOCX em %s", output_path)

    doc = Document()
    apply_abnt_styles(doc)
    chart_generator = ChartGenerator()
    fig_n = 1
    
    # Cabeçalho / Título
    doc.add_heading(TITULO_SECRETARIA.upper(), 0)
    
    p = doc.add_paragraph()
    p.add_run(f"Relatório Socioeconômico – {MUNICIPIO}/{UF}").bold = True
    p.add_run(f"\nPeríodo de Análise: {ano_inicio} a {ano_fim}")
    p.add_run(f"\nData de Emissão: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    # 1. Nota Metodológica (PIB)
    doc.add_heading("1. Nota Metodológica (Estimativa do PIB)", level=1)
    doc.add_paragraph(
        "A estimativa do Produto Interno Bruto Municipal foi elaborada a partir de metodologia híbrida, "
        "combinando o último dado oficial publicado pelo IBGE com proxies econômicos atualizados, tais como "
        "Valor Adicionado Fiscal (SEFAZ-MG) e Massa Salarial (RAIS/CAGED). A projeção foi "
        "realizada por meio de modelos de séries temporais, garantindo coerência estatística e aderência à dinâmica econômica local."
    )

    # 2. Dados de Projeção do PIB
    df_prev = get_estimativa_stored()
    if not df_prev.empty:
        doc.add_heading("2. Projeções de Curto Prazo", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Ano'
        hdr_cells[1].text = 'Valor Estimado (R$ mil)'
        hdr_cells[2].text = 'Status'
        
        for _, row in df_prev.sort_values("Ano").iterrows():
            try:
                row_cells = table.add_row().cells
                row_cells[0].text = str(int(row["Ano"]))
                row_cells[1].text = _fmt_num(row["Valor"])
                row_cells[2].text = "Projeção"
            except: continue

    # 3. Indicadores Socioeconômicos (Executivo)
    doc.add_heading("3. Análise Detalhada por Indicador", level=1)
    
    # Organizar indicadores por fonte para um relatório mais estruturado
    all_indicators = list_indicators()
    education_keys = {
        "MATRICULAS_TOTAL",
        "ESCOLAS_FUNDAMENTAL",
        "IDEB_ANOS_INICIAIS",
        "IDEB_ANOS_FINAIS",
        "TAXA_APROVACAO_FUNDAMENTAL",
    }
    
    current_source = None
    for ind in all_indicators:
        key = ind["indicator_key"]
        source = ind["source"]
        unit_catalog = (CATALOGO_INDICADORES.get(key) or {}).get("unidade", "")  # fallback amigável
        
        if key in education_keys and not str(source).startswith("INEP"):
            continue

        if source != current_source:
            current_source = source
            doc.add_heading(f"Fonte: {source}", level=1)

        df = get_timeseries(key, source=source)
        if df.empty: continue
        
        df = df[(df["Ano"] >= ano_inicio) & (df["Ano"] <= ano_fim)]
        if df.empty: continue
        
        # Nome amigável quando existir
        nome = (CATALOGO_INDICADORES.get(key) or {}).get("nome", key)
        unidade_df = ""
        try:
            unidade_df = str(df["Unidade"].dropna().iloc[-1]).strip()
        except Exception:
            unidade_df = ""
        unidade = unidade_df or unit_catalog

        doc.add_heading(f"{nome}", level=2)
        doc.add_paragraph(f"Chave: {key} | Fonte: {source}" + (f" | Unidade: {unidade}" if unidade else ""))
        
        # Análise de Tendência
        resultado = analisar_tendencia(
            df,
            ano_inicio=ano_inicio,
            ano_fim=ano_fim,
            text_format="plain",
            unidade=unidade or None,
        )
        doc.add_paragraph(f"Análise de Tendência: {resultado.resumo}").bold = True
        doc.add_paragraph(resultado.tendencia_texto)

        # Gráfico (Figura) – somente quando há série histórica suficiente
        try:
            if len(df.dropna(subset=["Ano", "Valor"])) >= 2:
                charts_dir = (DATA_DIR / "charts").resolve()
                charts_dir.mkdir(parents=True, exist_ok=True)
                chart_path = charts_dir / f"{key}.png"
                ylab = f"Valor ({unidade})" if unidade else "Valor"
                created = chart_generator.create_line_chart(
                    df, "Ano", "Valor", f"{nome} – Série histórica", xlabel="Ano", ylabel=ylab, output_path=chart_path
                )
                if created and created.exists():
                    doc.add_picture(str(created), width=Inches(6.0))
                    add_figure_caption(
                        doc,
                        figure_number=fig_n,
                        title=f"{nome} – Série histórica" + (f" ({unidade})" if unidade else ""),
                        source=source,
                    )
                    fig_n += 1
        except Exception as e:
            logger.warning("Falha ao gerar gráfico para %s: %s", key, e)
        
        # Tabela de Dados Gerais
        doc.add_paragraph("Série Histórica Selecionada:")
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Ano'
        hdr_cells[1].text = f"Valor ({unidade})" if unidade else "Valor"
        
        for _, row in df.sort_values("Ano").iterrows():
            try:
                row_cells = table.add_row().cells
                row_cells[0].text = str(int(row["Ano"]))
                row_cells[1].text = _fmt_num(row["Valor"])
            except: continue
            
        doc.add_paragraph("\n") # Espaçador

    # Rodapé automático (python-docx lida via seções mas simplificaremos)
    doc.add_paragraph("-" * 30)
    doc.add_paragraph(f"Observatório Socioeconômico de {MUNICIPIO} - Documento gerado automaticamente.")

    doc.save(str(output_path))
    return output_path
