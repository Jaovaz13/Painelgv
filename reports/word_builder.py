"""
Construtor de relatórios Word profissional com estrutura institucional completa.
Implementa todos os blocos temáticos com análise inteligente e design profissional.
"""

import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

from analytics.tendencias import analisar_tendencia
from analytics.estimativa_pib import get_estimativa_stored
from config import MUNICIPIO, UF, DATA_DIR
from database import get_timeseries, list_indicators
from reports.abnt import apply_abnt_styles, add_figure_caption
from reports.structure import ReportSection, create_empty_report_structure, BRAND_COLORS
from reports.text_engine import TextGenerator, TrendAnalyzer, analyze_multiple_indicators
from reports.charts import ChartGenerator, create_thematic_charts
from reports.indicator_groups import INDICATOR_GROUPS, organize_indicators_by_groups, clean_indicators_list
from reports.text_enhancer import text_enhancer
from reports.executive_summary import executive_summary_generator

logger = logging.getLogger(__name__)

TITULO_SECRETARIA = "Secretaria Municipal de Desenvolvimento, Ciência, Tecnologia e Inovação"

class WordReportBuilder:
    """Construtor profissional de relatórios Word com estrutura institucional."""
    
    def __init__(self):
        """Inicializa o construtor com configurações padrão."""
        self.doc = None
        self.chart_generator = ChartGenerator()
        self.text_generator = TextGenerator()
        self.charts_dir = DATA_DIR / "charts"
        self.charts_dir.mkdir(exist_ok=True)
        self._figure_counter = 1
    
    def _add_custom_heading(self, text: str, level: int, color: str = None):
        """Adiciona título customizado com cor e formatação."""
        heading = self.doc.add_heading(text, level)
        
        # Aplicar cor customizada se especificada
        if color and color in BRAND_COLORS:
            for run in heading.runs:
                r = run._element
                rPr = r.get_or_add_rPr()
                color_elem = OxmlElement('w:color')
                color_elem.set(qn('w:val'), color.lstrip('#'))
                rPr.append(color_elem)
        
        return heading
    
    def _add_section_break(self):
        """Adiciona quebra de seção."""
        # python-docx não possui WD_BREAK.SECTION; usar nova seção em nova página.
        self.doc.add_section(WD_SECTION.NEW_PAGE)
    
    def _format_paragraph(self, text: str, bold: bool = False, italic: bool = False, 
                         size: int = None, color: str = None):
        """Formata parágrafo com estilo específico."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if size:
            run.font.size = Pt(size)
        if color and color in BRAND_COLORS:
            run.font.color.rgb = RGBColor(*[int(color.lstrip('#')[i:i+2], 16) for i in (0, 2, 4)])
        
        return p
    
    def _create_styled_table(self, headers: List[str], data: List[List[Any]], 
                            style: str = 'Table Grid') -> Any:
        """Cria tabela com estilo profissional."""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Cabeçalho
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # Negrito no cabeçalho
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Dados
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        
        return table
    
    def _add_chart_image(self, chart_path: Optional[Path], width: float = 6.0):
        """Adiciona imagem de gráfico ao documento."""
        if chart_path and chart_path.exists():
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(chart_path), width=Inches(width))
        elif chart_path and hasattr(chart_path, 'getvalue'):
            # Handle BytesIO object
            try:
                # Salvar BytesIO temporariamente
                temp_path = self.charts_dir / f"temp_chart_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
                with open(temp_path, 'wb') as f:
                    f.write(chart_path.getvalue())
                
                # Adicionar imagem do arquivo temporário
                paragraph = self.doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(temp_path), width=Inches(width))
                
                # Limpar arquivo temporário
                temp_path.unlink()
                
            except Exception as e:
                logger.error(f"Erro ao processar BytesIO: {e}")
                self._format_paragraph("*Gráfico não disponível*")
        else:
            # Se não houver gráfico, adicionar nota
            self._format_paragraph("*Gráfico não disponível para este indicador*")
    
    def _build_institutional_block(self, ano_inicio: int, ano_fim: int):
        """Constrói Bloco A - Institucional."""
        self._add_custom_heading("BLOCO A – INSTITUCIONAL", 1, BRAND_COLORS["primary"])
        
        # Capa
        self._add_custom_heading("Apresentação Institucional", 2)
        self._format_paragraph(f"{TITULO_SECRETARIA}", bold=True, size=14)
        self._format_paragraph(f"Município de {MUNICIPIO} – {UF}", bold=True, size=12)
        self._format_paragraph(f"Relatório Socioeconômico Municipal", bold=True, size=13)
        self._format_paragraph(f"Período de Análise: {ano_inicio} a {ano_fim}")
        self._format_paragraph(f"Data de Emissão: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        
        self._add_section_break()
        
        # Escopo e Fontes
        self._add_custom_heading("Escopo e Fontes de Dados", 2)
        self._format_paragraph("Este relatório apresenta análise integrada dos principais indicadores socioeconômicos do município, "
                              "com base em dados oficiais e atualizados automaticamente das seguintes fontes:")
        
        fontes = [
            "• IBGE/SIDRA: População, PIB e indicadores demográficos",
            "• SEFAZ-MG: Valor Adicionado Fiscal e capacidade fiscal", 
            "• SEBRAE: Empreendedorismo e empresas ativas",
            "• CAGED/RAIS: Mercado de trabalho formal",
            "• INEP: Educação e matrículas escolares",
            "• DataSUS: Saúde e indicadores de mortalidade",
            "• SEEG: Emissões de gases de efeito estufa",
            "• MapBiomas: Uso do solo e cobertura vegetal",
            "• IDSC-BR: Índice de desenvolvimento sustentável"
        ]
        
        for fonte in fontes:
            self._format_paragraph(fonte)
        
        self._add_section_break()
    
    def _build_executive_block(self, indicators_data: Dict[str, pd.DataFrame], 
                              ano_inicio: int, ano_fim: int):
        """Constrói Bloco B – Executivo com resumo estratégico."""
        self._add_custom_heading("BLOCO B – EXECUTIVO", 1, BRAND_COLORS["primary"])
        
        # Resumo Executivo Estratégico (NOVO)
        self._add_custom_heading("Resumo Executivo Estratégico", 2)
        
        # Gerar resumo executivo melhorado
        executive_summary = executive_summary_generator.generate_executive_summary(
            indicators_data, ano_inicio, ano_fim
        )
        self._format_paragraph(executive_summary)
        
        self._add_section_break()
        
        # Destaques Principais
        self._add_custom_heading("Principais Destaques do Período", 2)
        
        # Analisar todos os indicadores
        all_analyses = analyze_multiple_indicators(indicators_data)
        
        # Destaques positivos
        highlights = []
        for name, analysis in all_analyses.items():
            if analysis.get("strength", 0) > 60 and analysis.get("direction") == "increasing":
                df = indicators_data.get(name)
                if df is not None and not df.empty:
                    latest = df.iloc[-1]
                    # Melhorar interpretação
                    interpretation = text_enhancer.enhance_text(analysis.get("interpretation", ""))
                    highlights.append((name, latest, interpretation))
        
        if highlights:
            def _get_strength(item):
                _, _, interp = item
                return interp.get("strength", 0) if isinstance(interp, dict) else 0

            highlights.sort(key=_get_strength, reverse=True)

            for name, latest, interpretation in highlights[:5]:
                value = latest.get("Valor")
                year = int(latest.get("Ano"))
                unit = str(latest.get("Unidade", "")).strip()
                unidade_textual = f" {unit}" if unit else ""
                # Sem emojis/markdown no documento institucional
                self._format_paragraph(f"{name}: {value:,.0f}{unidade_textual} em {year}", bold=True)
                self._format_paragraph(str(interpretation))
                self._format_paragraph("")  # Espaçamento
        
        self._add_section_break()
        
        # Pontos de Atenção
        self._add_custom_heading("Pontos de Atenção Prioritários", 2)
        
        attention_indicators = []
        for name, analysis in all_analyses.items():
            if analysis.get("direction") == "decreasing" and analysis.get("strength", 0) > 40:
                df = indicators_data.get(name)
                if df is not None and not df.empty:
                    latest = df.iloc[-1]
                    # Melhorar interpretação para quedas
                    interpretation = text_enhancer.enhance_strong_decline(
                        analysis.get("total_variation", 0), name
                    )
                    attention_indicators.append((name, latest, interpretation))
        
        if attention_indicators:
            for name, latest, interpretation in attention_indicators[:5]:
                value = latest.get("Valor", latest.get("value"))
                year = int(latest.get("Ano", latest.get("year")))
                unit = str(latest.get("Unidade", "")).strip()
                unidade_textual = f" {unit}" if unit else ""
                self._format_paragraph(f"{name}: {value:,.0f}{unidade_textual} em {year}", bold=True)
                self._format_paragraph(str(interpretation))
                self._format_paragraph("")  # Espaçamento
        else:
            self._format_paragraph("✅ Não foram identificados pontos críticos no período analisado.")
        
        self._add_section_break()
    
    def _build_thematic_blocks_by_groups(self, indicators_data: Dict[str, pd.DataFrame]):
        """Constrói Bloco C com análise temática organizada por grupos."""
        organized_indicators = organize_indicators_by_groups(list(indicators_data.keys()))

        self._add_custom_heading("BLOCO C – ANÁLISE TEMÁTICA", 1, BRAND_COLORS["primary"])

        for group_name, indicator_list in organized_indicators.items():
            if not indicator_list:
                continue

            group_config = INDICATOR_GROUPS.get(group_name, {})
            group_title = group_config.get("title", group_name.title())
            group_description = group_config.get("description", "")

            self._add_custom_heading(group_title, 2)
            if group_description:
                self._format_paragraph(group_description)

            for indicator in indicator_list:
                if indicator in indicators_data:
                    self._build_indicator_analysis(indicator, indicators_data[indicator])

        self._add_section_break()
    
    def _build_indicator_analysis(self, indicator_name: str, data: pd.DataFrame):
        """Constrói análise individual de um indicador."""
        # Título do indicador
        self._add_custom_heading(f"Análise: {indicator_name}", 2)
        
        # Tabela de dados
        # Cabeçalho exibe unidade quando disponível
        unit = None
        if "Unidade" in data.columns:
            try:
                unit = str(data["Unidade"].iloc[-1]).strip()
            except Exception:
                unit = None

        valor_header = "Valor"
        if unit:
            valor_header = f"Valor ({unit})"

        headers = ["Ano", valor_header, "Variação %"]
        table_data = []
        
        # Garantir ordem temporal e índice posicional
        data_sorted = data.sort_values("Ano").reset_index(drop=True)

        for pos in range(len(data_sorted)):
            row = data_sorted.iloc[pos]
            year = int(row["Ano"])
            value = row["Valor"]

            # Calcular variação percentual (posicional)
            if pos > 0:
                prev_value = data_sorted.iloc[pos - 1]["Valor"]
                if prev_value != 0:
                    var_pct = ((value - prev_value) / prev_value) * 100
                    var_str = f"{var_pct:+.1f}%"
                else:
                    var_str = "N/A"
            else:
                var_str = "—"

            table_data.append([year, f"{value:,.2f}", var_str])
        
        # Inverter ordem para mostrar mais recente primeiro
        table_data = table_data[::-1]
        self._create_styled_table(headers, table_data, 'Table Grid')
        
        # Análise de tendência
        try:
            analysis = TrendAnalyzer.analyze_trend(data["Valor"])
            
            # Melhorar interpretação
            enhanced_interpretation = text_enhancer.enhance_text(analysis.get("interpretation", ""))
            
            # Formatar análise (sem markdown)
            self._format_paragraph("Análise de tendência:", bold=True)
            self._format_paragraph(f"Direção: {analysis.get('direction', 'estável').title()}")
            self._format_paragraph(f"Força: {analysis.get('strength', 0):.2f}")
            self._format_paragraph(f"Confiança: {analysis.get('confidence', 0):.2f}")
            self._format_paragraph(f"Interpretação: {enhanced_interpretation}")
            
            # Adicionar gráfico
            chart_path = self.charts_dir / f"{indicator_name}_trend.png"
            self.chart_generator.create_line_chart(
                data, "Ano", "Valor", f"Evolução de {indicator_name}",
                output_path=str(chart_path)
            )
            self._add_chart_image(chart_path, width=6.0)

            # Legenda ABNT da figura (quando houver gráfico gerado)
            if chart_path.exists():
                unit = ""
                try:
                    unit = str(data["Unidade"].dropna().iloc[-1]).strip()
                except Exception:
                    unit = ""
                src = ""
                try:
                    src = str(data["source"].dropna().iloc[-1]).strip()
                except Exception:
                    src = ""
                fig_title = f"Evolução de {indicator_name}" + (f" ({unit})" if unit else "")
                add_figure_caption(
                    self.doc,
                    figure_number=self._figure_counter,
                    title=fig_title,
                    source=src or "Base de Dados Integrada (Painel GV)",
                )
                self._figure_counter += 1
            
        except Exception as e:
            logger.error(f"Erro na análise de {indicator_name}: {e}")
            self._format_paragraph("*Dados insuficientes para análise de tendência*")
        
        self._format_paragraph("")  # Espaçamento
    
    def _build_thematic_block(self, theme: str, indicators_data: Dict[str, pd.DataFrame]):
        self._add_custom_heading(f"BLOCO C – ANÁLISE TEMÁTICA: {theme.upper()}", 1, BRAND_COLORS["primary"])
        
        # Filtrar indicadores do tema
        theme_indicators = self._filter_indicators_by_theme(theme, indicators_data)
        
        if not theme_indicators:
            self._format_paragraph(f"Dados insuficientes para análise do tema {theme}.")
            return
        
        # Contextualização
        self._add_custom_heading("Contextualização", 2)
        context = self._get_theme_context(theme)
        self._format_paragraph(context)
        
        self._add_section_break()
        
        # Tabela-síntese
        self._add_custom_heading("Tabela-Síntese", 2)
        self._create_theme_summary_table(theme_indicators)
        
        self._add_section_break()
        
        # Gráficos históricos
        self._add_custom_heading("Evolução Histórica", 2)
        
        # Gerar gráficos do tema
        theme_charts = create_thematic_charts(theme_indicators, self.charts_dir)
        
        for indicator_name, chart_path in theme_charts.items():
            self._format_paragraph(f"**{indicator_name}**", bold=True)
            self._add_chart_image(chart_path, width=6.0)
            self._format_paragraph("")  # Espaçamento
        
        self._add_section_break()
        
        # Análise automática de tendência
        self._add_custom_heading("Análise de Tendência", 2)
        
        theme_analysis = self.text_generator.generate_thematic_analysis(theme, theme_indicators)
        # Adicionar apenas a análise de tendência (sem markdown)
        lines = theme_analysis.split('\n')
        for line in lines:
            if line.strip() and not line.startswith('#'):
                self._format_paragraph(line.strip())
        
        self._add_section_break()
    
    def _filter_indicators_by_theme(self, theme: str, indicators_data: Dict[str, pd.DataFrame]) -> Dict[str, pd.DataFrame]:
        """Filtra indicadores por tema específico."""
        theme_keywords = {
            "economia": ["pib", "vaf", "empresa", "receita"],
            "trabalho_renda": ["emprego", "salario", "caged", "rais", "trabalho"],
            "educacao": ["matricula", "escola", "educacao", "inep"],
            "saude": ["mortalidade", "saude", "datasus", "obito"],
            "sustentabilidade": ["idsc", "emissao", "area", "mapbiomas", "seeg", "vegetacao", "urbana"]
        }
        
        keywords = theme_keywords.get(theme.lower(), [])
        filtered = {}
        
        for name, df in indicators_data.items():
            if any(keyword in name.lower() for keyword in keywords):
                filtered[name] = df
        
        return filtered
    
    def _get_theme_context(self, theme: str) -> str:
        """Retorna contextualização específica do tema."""
        contexts = {
            "economia": f"A análise econômica de {MUNICIPIO} avalia a dinâmica produtiva, a capacidade de geração de riqueza e o ambiente de negócios do município, considerando o PIB, a estrutura produtiva e a capacidade fiscal.",
            "trabalho_renda": f"O mercado de trabalho formal e os indicadores de renda em {MUNICIPIO} refletem a capacidade de geração de empregos, a massa salarial e as oportunidades de empreendedorismo na região.",
            "educacao": f"Os indicadores educacionais de {MUNICIPIO} medem o acesso ao ensino, a qualidade da educação e o desenvolvimento de capital humano, fundamentais para o desenvolvimento sustentável.",
            "saude": f"Os indicadores de saúde em {MUNICIPIO} avaliam o acesso aos serviços, os resultados de saúde pública e o bem-estar da população, refletindo a qualidade de vida local.",
            "sustentabilidade": f"A análise de sustentabilidade em {MUNICIPIO} integra indicadores ambientais, sociais e de desenvolvimento sustentável, essenciais para o planejamento de longo prazo."
        }
        return contexts.get(theme, f"Análise temática específica para {theme}.")
    
    def _create_theme_summary_table(self, theme_indicators: Dict[str, pd.DataFrame]):
        """Cria tabela-síntese do tema."""
        headers = ["Indicador", "Último Valor", "Ano", "Tendência", "Força"]
        data = []
        
        for name, df in theme_indicators.items():
            if df.empty:
                continue
            
            latest = df.iloc[-1]
            analysis = TrendAnalyzer.analyze_trend(df["Valor"])
            
            direction_map = {
                "increasing": "📈 Crescimento",
                "decreasing": "📉 Queda", 
                "stable": "➡️ Estável",
                "insufficient_data": "❓ Dados insuf."
            }
            
            trend = direction_map.get(analysis["direction"], "❓")
            strength = f"{analysis['strength']:.0f}%" if analysis['strength'] > 0 else "N/A"
            
            unidade = str(latest.get("Unidade", "")).strip()
            valor_fmt = f"{latest['Valor']:,.0f}"
            if unidade:
                valor_fmt = f"{valor_fmt} {unidade}"

            data.append(
                [
                    name,
                    valor_fmt,
                    int(latest["Ano"]),
                    trend,
                    strength,
                ]
            )
        
        if data:
            self._create_styled_table(headers, data, 'Light Shading Accent 1')
    
    def _build_comparisons_block(self, indicators_data: Dict[str, pd.DataFrame]):
        """Constrói Bloco D - Comparações e Tendências."""
        self._add_custom_heading("BLOCO D – COMPARAÇÕES E TENDÊNCIAS", 1, BRAND_COLORS["primary"])
        
        # Evolução temporal
        self._add_custom_heading("Evolução Temporal Comparativa", 2)
        
        # Selecionar indicadores principais para comparação
        main_indicators = {}
        for name, df in indicators_data.items():
            if any(keyword in name.lower() for keyword in ["pib", "emprego", "empresa", "idsc"]):
                if not df.empty and len(df) > 2:
                    main_indicators[name] = df
        
        if main_indicators:
            # Criar gráfico comparativo
            chart_path = self.charts_dir / "comparativo_eixos.png"
            self.chart_generator.create_comparison_chart(
                main_indicators, "Ano", "Valor", 
                "Evolução Comparativa dos Principais Indicadores",
                output_path=chart_path
            )
            self._add_chart_image(chart_path, width=7.0)
        
        self._add_section_break()
        
        # Variação percentual
        self._add_custom_heading("Variação Percentual", 2)
        
        variations = []
        for name, df in main_indicators.items():
            if len(df) >= 2:
                first = df.iloc[0]["Valor"]
                last = df.iloc[-1]["Valor"]
                var_pct = ((last - first) / first) * 100 if first != 0 else 0
                variations.append((name, var_pct, int(df.iloc[-1]["Ano"])))
        
        if variations:
            variations.sort(key=lambda x: abs(x[1]), reverse=True)
            
            headers = ["Indicador", "Variação %", "Ano Final"]
            data = []
            
            for name, var_pct, year in variations:
                var_symbol = "📈" if var_pct > 0 else "📉" if var_pct < 0 else "➡️"
                data.append([name, f"{var_symbol} {var_pct:+.1f}%", year])
            
            self._create_styled_table(headers, data, 'Light Shading Accent 2')
        
        self._add_section_break()
    
    def _build_sustainability_block(self, indicators_data: Dict[str, pd.DataFrame]):
        """Constrói Bloco E - Sustentabilidade e Desenvolvimento."""
        self._add_custom_heading("BLOCO E – SUSTENTABILIDADE E DESENVOLVIMENTO", 1, BRAND_COLORS["primary"])
        
        # Índices compostos
        self._add_custom_heading("Índices Compostos", 2)
        
        idsc_data = indicators_data.get("IDSC_GERAL")
        if idsc_data is not None and not idsc_data.empty:
            latest = idsc_data.iloc[-1]
            analysis = TrendAnalyzer.analyze_trend(idsc_data["Valor"])
            
            unidade = str(latest.get("Unidade", "")).strip()
            unidade_textual = f" {unidade}" if unidade else ""
            self._format_paragraph(
                f"IDSC-BR (score geral): {latest['Valor']:.2f}{unidade_textual} em {int(latest['Ano'])}",
                bold=True,
            )
            self._format_paragraph(analysis["interpretation"])
            
            # Gráfico do IDSC
            chart_path = self.charts_dir / "idsc_evolution.png"
            self.chart_generator.create_line_chart(
                idsc_data, "Ano", "Valor", "Evolução do IDSC-BR",
                output_path=chart_path
            )
            self._add_chart_image(chart_path, width=6.0)
        
        self._add_section_break()
        
        # Evolução integrada
        self._add_custom_heading("Evolução Integrada", 2)
        
        sustainability_indicators = {}
        for name, df in indicators_data.items():
            if any(keyword in name.lower() for keyword in ["emissao", "area", "vegetacao", "urbana"]):
                if not df.empty:
                    sustainability_indicators[name] = df
        
        if sustainability_indicators:
            chart_path = self.charts_dir / "sustainability_integrated.png"
            self.chart_generator.create_comparison_chart(
                sustainability_indicators, "Ano", "Valor",
                "Evolução dos Indicadores de Sustentabilidade",
                output_path=chart_path
            )
            self._add_chart_image(chart_path, width=7.0)
        
        self._add_section_break()
    
    def _build_conclusions_block(self, indicators_data: Dict[str, pd.DataFrame]):
        """Constrói Bloco F - Conclusões Estratégicas."""
        self._add_custom_heading("BLOCO F – CONCLUSÕES ESTRATÉGICAS", 1, BRAND_COLORS["primary"])
        
        # Análise completa
        all_analyses = analyze_multiple_indicators(indicators_data)
        
        # Síntese automática
        self._add_custom_heading("Síntese Automática", 2)
        conclusions_text = self.text_generator.generate_strategic_conclusions(all_analyses)
        
        # Adicionar texto sem markdown
        lines = conclusions_text.split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('##'):
                current_section = line.replace('##', '').strip()
                self._add_custom_heading(current_section, 2)
            elif line.startswith('###'):
                subtitle = line.replace('###', '').strip()
                self._add_custom_heading(subtitle, 3)
            elif line.startswith('•'):
                self._format_paragraph(line)
            elif line.startswith(('1.', '2.', '3.', '4.', '5.')):
                self._format_paragraph(line, bold=True)
            else:
                if current_section and line != current_section:
                    self._format_paragraph(line)
        
        self._add_section_break()
    
    def _build_methodology_block(self):
        """Constrói Bloco G - Metodologia e Transparência."""
        self._add_custom_heading("BLOCO G – METODOLOGIA E TRANSPARÊNCIA", 1, BRAND_COLORS["primary"])
        
        # Fontes detalhadas
        self._add_custom_heading("Fontes de Dados", 2)
        
        fontes_detalhadas = [
            ("IBGE/SIDRA", "https://apisidra.ibge.gov.br/", "API", "População, PIB, indicadores demográficos"),
            ("SEFAZ-MG", "https://www.fazenda.mg.gov.br/empresas/vaf/", "CSV", "Valor Adicionado Fiscal"),
            ("SEBRAE", "https://datasebrae.com.br/municipios/", "CSV", "Empreendedorismo, empresas"),
            ("CAGED/RAIS", "https://www.gov.br/trabalho-e-emprego/pt-br/assuntos/estatisticas", "XLSX", "Mercado de trabalho"),
            ("INEP", "https://www.gov.br/inep/pt-br/acesso-a-informacao/dados-abertos", "XLSX", "Educação"),
            ("DataSUS", "https://datasus.saude.gov.br/informacoes-de-saude-tabnet/", "CSV", "Saúde"),
            ("SEEG", "https://seeg.eco.br/dados/", "CSV", "Emissões GEE"),
            ("MapBiomas", "https://mapbiomas.org/download", "XLSX", "Uso do solo"),
            ("IDSC-BR", "https://idsc.cidadessustentaveis.org.br", "XLSX", "Desenvolvimento sustentável")
        ]
        
        headers = ["Fonte", "URL", "Formato", "Indicadores"]
        data = [[fonte, url, formato, indicadores] for fonte, url, formato, indicadores in fontes_detalhadas]
        self._create_styled_table(headers, data, 'Table Grid')
        
        self._add_section_break()
        
        # Métodos de análise
        self._add_custom_heading("Métodos de Análise", 2)
        
        metodos = [
            "• Análise de tendência: regressão linear com coeficiente R² e teste de significância",
            "• Projeções: modelos de séries temporais com intervalos de confiança",
            "• Comparação: análise percentual e correlação entre indicadores",
            "• Síntese: geração automática de insights baseada em padrões estatísticos",
            "• Validação: verificação de consistência e qualidade dos dados",
        ]
        
        for metodo in metodos:
            self._format_paragraph(metodo)
        
        self._add_section_break()
        
        # Atualização
        self._add_custom_heading("Atualização dos Dados", 2)
        self._format_paragraph(f"Data de extração: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        self._format_paragraph("Frequência de atualização: Conforme disponibilidade das fontes oficiais")
        self._format_paragraph("Processo: Coleta automática com validação de qualidade")
        
        self._add_section_break()
        
        # Limitações
        self._add_custom_heading("Limitações Metodológicas", 2)
        
        limitacoes = [
            "• Disponibilidade de dados pode variar entre fontes",
            "• Indicadores com séries históricas curtas têm menor confiabilidade estatística",
            "• Projeções baseiam-se em tendências históricas e podem não capturar mudanças estruturais",
            "• Análises automáticas requerem validação por especialistas do domínio"
        ]
        
        for limitacao in limitacoes:
            self._format_paragraph(limitacao)
        
        self._add_section_break()
    
    def build_complete_report(self, ano_inicio: int, ano_fim: int, 
                           output_path: Optional[str] = None) -> Path:
        """
        Constrói relatório completo com todos os blocos institucionais.
        
        Args:
            ano_inicio: Ano inicial da análise
            ano_fim: Ano final da análise
            output_path: Caminho para salvar o relatório
            
        Returns:
            Path do arquivo gerado
        """
        output_path = Path(output_path) if output_path else DATA_DIR / f"relatorio_institucional_{ano_inicio}_{ano_fim}.docx"
        output_path = output_path.resolve()
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        logger.info("Iniciando construção do relatório institucional completo")
        
        # Inicializar documento
        self.doc = Document()
        apply_abnt_styles(self.doc)
        
        # Coletar dados dos indicadores
        indicators_data = self._collect_indicators_data(ano_inicio, ano_fim)
        
        # Construir blocos
        self._build_institutional_block(ano_inicio, ano_fim)
        self._build_executive_block(indicators_data, ano_inicio, ano_fim)
        self._build_thematic_blocks_by_groups(indicators_data)
        self._build_comparisons_block(indicators_data)
        self._build_conclusions_block(indicators_data)
        self._build_methodology_block()
        
        # Rodapé final
        self._add_section_break()
        self._format_paragraph("=" * 50)
        self._format_paragraph(f"Observatório Socioeconômico de {MUNICIPIO} - Documento gerado automaticamente", 
                              italic=True, size=10)
        self._format_paragraph(f"Endereço eletrônico: observatorio@{MUNICIPIO}.mg.gov.br", italic=True, size=10)
        
        # Salvar documento
        self.doc.save(str(output_path))
        logger.info(f"Relatório institucional completo gerado em {output_path}")
        
        return output_path
    
    def _collect_indicators_data(self, ano_inicio: int, ano_fim: int) -> Dict[str, pd.DataFrame]:
        """Coleta dados de todos os indicadores para análise."""
        indicators_data = {}
        all_indicators = list_indicators()
        education_keys = {
            "MATRICULAS_TOTAL",
            "ESCOLAS_FUNDAMENTAL",
            "IDEB_ANOS_INICIAIS",
            "IDEB_ANOS_FINAIS",
            "TAXA_APROVACAO_FUNDAMENTAL",
        }
        
        for ind in all_indicators:
            key = ind["indicator_key"]
            source = ind["source"]

            # EDUCAÇÃO: usar exclusivamente dados provenientes de arquivos reais em data/raw
            if key in education_keys and source != "INEP_RAW":
                continue
            
            try:
                df = get_timeseries(key, source=source)
                if not df.empty:
                    # Filtrar por período
                    df_filtered = df[(df["Ano"] >= ano_inicio) & (df["Ano"] <= ano_fim)]
                    if not df_filtered.empty:
                        indicators_data[key] = df_filtered
            except Exception as e:
                logger.warning(f"Erro ao coletar indicador {key}: {e}")
                continue
        
        return indicators_data

# Função de conveniência para manter compatibilidade
def gerar_relatorio_docx(
    ano_inicio: int,
    ano_fim: int,
    output_path: str | Path | None = None,
) -> Path:
    """
    Gera relatório Word institucional completo (função de compatibilidade).
    
    Args:
        ano_inicio: Ano inicial da análise
        ano_fim: Ano final da análise
        output_path: Caminho para salvar o relatório
        
    Returns:
        Path do arquivo gerado
    """
    builder = WordReportBuilder()
    return builder.build_complete_report(ano_inicio, ano_fim, output_path)
